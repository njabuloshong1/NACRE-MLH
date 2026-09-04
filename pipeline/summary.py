"""Cross-run tables and a Word report.

Every table is also written as CSV, so the report is a convenience rather than the only copy of a
number. If python-docx is not installed the CSVs and a Markdown summary are still produced.
"""
import os, warnings; warnings.filterwarnings("ignore")
import numpy as np, pandas as pd

LEVELS = ["subtype", "lineage", "compartment"]
TIERS = LEVELS + ["unresolved"]
PLATFORM_NAME = {"X": "Xenium", "M": "MERSCOPE", "C": "CosMx"}


def _f(x, n=2):
    return "" if x is None or (isinstance(x, float) and np.isnan(x)) else f"{x:.{n}f}"


def build_tables(runs):
    """runs: list of dicts from run_pipeline, each carrying name/platform/cells/conc/res/acc."""
    T = {}
    T["datasets"] = pd.DataFrame([{
        "Run": r["name"], "Platform": PLATFORM_NAME.get(r["platform"], r["platform"]),
        "Cells": r.get("cells"), "Genes": r.get("genes"),
        "Reference cells": r.get("ref_cells"), "Reference types": r.get("ref_types"),
        "Panel genes shared": r.get("shared"),
        "Median transcripts": r.get("median_tx"),
        "Runtime (min)": r.get("minutes"),
    } for r in runs])

    rows = []
    for r in runs:
        c = r.get("conc")
        if c is None: continue
        d = {"Run": r["name"]}
        for lev in LEVELS:
            s = c[c.level == lev]
            if len(s):
                d[f"{lev} before"] = float(s.before.iloc[0])
                d[f"{lev} after"] = float(s.after.iloc[0])
                d[f"{lev} delta"] = float(s.delta.iloc[0])
        rows.append(d)
    T["concordance"] = pd.DataFrame(rows)

    T["certificate"] = pd.DataFrame([dict(
        {"Run": r["name"], "Cells": r.get("cells")},
        **{t: r.get("res", {}).get(t) for t in TIERS},
        **{f"{t} %": (100 * r["res"][t] / r["cells"]) if r.get("res") and r.get("cells") else None
           for t in TIERS})
        for r in runs])

    T["votes"] = pd.DataFrame([{
        "Run": r["name"], "LOCK %": r.get("lock"), "STRONG %": r.get("strong"),
        "SPLIT %": r.get("split")} for r in runs])

    acc = [r for r in runs if r.get("acc")]
    if acc:
        T["accuracy"] = pd.DataFrame([dict({"Run": r["name"]}, **r["acc"]) for r in acc])
    return {k: v for k, v in T.items() if len(v)}


def write_csvs(T, out):
    os.makedirs(out, exist_ok=True)
    for k, df in T.items():
        df.to_csv(os.path.join(out, f"summary_{k}.csv"), index=False)


def _md_table(df):
    """Written out by hand rather than via DataFrame.to_markdown, which needs `tabulate`; the
    summary should not fail for want of an optional formatting package."""
    def cell(v):
        if v is None or (isinstance(v, float) and np.isnan(v)): return ""
        if isinstance(v, float): return f"{v:,.2f}"
        if isinstance(v, (int, np.integer)): return f"{v:,}"
        return str(v).replace("|", r"\|")
    cols = [str(c) for c in df.columns]
    out = ["| " + " | ".join(cols) + " |", "|" + "|".join("---" for _ in cols) + "|"]
    for _, row in df.iterrows():
        out.append("| " + " | ".join(cell(v) for v in row) + " |")
    return "\n".join(out)


def write_markdown(T, out, title):
    p = os.path.join(out, "SUMMARY.md")
    with open(p, "w", encoding="utf-8") as fh:
        fh.write(f"# {title}\n\n")
        for k, df in T.items():
            fh.write(f"## {k.capitalize()}\n\n{_md_table(df)}\n\n")
    return p


def write_docx(T, out, title, figures=None):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10)
    doc.add_heading(title, 0)
    doc.add_paragraph(
        "Consensus annotation with a resolution certificate. Each cell is labelled by four "
        "independent annotators; a learned resolver settles contested cells, and the certificate "
        "records the deepest level (subtype, lineage or compartment) at which the annotators "
        "actually agree. Cells with no agreement at any level are reported as unresolved rather "
        "than given a label the evidence does not support.")

    CAPTION = {
        "datasets": "Datasets processed, with reference and panel details.",
        "concordance": "Mean pairwise concordance among the four annotators, before and after the "
                       "resolver, at each level of the hierarchy.",
        "certificate": "Resolution certificate: how many cells reached each level.",
        "votes": "Vote structure: LOCK is 4/4 agreement, STRONG is 3/1, SPLIT is everything else.",
        "accuracy": "Agreement with the supplied truth column. Reported only for runs whose run "
                    "sheet named one.",
    }
    n = 0
    for k, df in T.items():
        n += 1
        doc.add_heading(f"Table {n}. {k.capitalize()}", level=2)
        doc.add_paragraph(CAPTION.get(k, ""))
        t = doc.add_table(rows=1, cols=len(df.columns)); t.style = "Light Grid Accent 1"
        for j, c in enumerate(df.columns):
            cell = t.rows[0].cells[j]; cell.text = str(c)
            for r in cell.paragraphs:
                for run in r.runs: run.font.bold = True; run.font.size = Pt(8.5)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for j, v in enumerate(row):
                cells[j].text = ("" if v is None or (isinstance(v, float) and np.isnan(v))
                                 else (f"{v:,.2f}" if isinstance(v, float)
                                       else f"{v:,}" if isinstance(v, (int, np.integer)) else str(v)))
                for r in cells[j].paragraphs:
                    for run in r.runs: run.font.size = Pt(8.5)

    if figures:
        doc.add_page_break()
        doc.add_heading("Figures", level=1)
        for run_name, made in figures.items():
            files = [(k, v) for k, v in made.items()
                     if isinstance(v, str) and v.lower().endswith(".png")]
            if not files: continue
            doc.add_heading(run_name, level=2)
            for key, fn in files:
                path = os.path.join(out, run_name, "figures", fn)
                if not os.path.exists(path): continue
                try:
                    doc.add_picture(path, width=Inches(6.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph(f"{run_name} — {key}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in cap.runs: r.font.size = Pt(8); r.font.italic = True
                except Exception:
                    continue

    p = os.path.join(out, "NACRE-MLH_report.docx")
    doc.save(p)
    return p
